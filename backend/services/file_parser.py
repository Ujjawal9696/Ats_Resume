"""
PDF/DOC Parser Service - Extract text from PDF, DOC, DOCX files
"""

import io
import os
from typing import Tuple
from pathlib import Path
from loguru import logger


def extract_text_from_pdf(file_content: bytes) -> Tuple[str, int]:
    """Extract text from PDF using pdfplumber with PyMuPDF fallback."""
    text = ""
    page_count = 0

    # Try pdfplumber first
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            logger.info(f"pdfplumber: extracted {len(text)} chars, {page_count} pages")
            return text.strip(), page_count
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")

    # Fallback: PyMuPDF (fitz)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_content, filetype="pdf")
        page_count = doc.page_count
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        logger.info(f"PyMuPDF: extracted {len(text)} chars, {page_count} pages")
        return text.strip(), page_count
    except Exception as e:
        logger.error(f"PyMuPDF also failed: {e}")
        raise ValueError("Could not extract text from PDF. File may be corrupted or image-based.")


def extract_text_from_docx(file_content: bytes) -> Tuple[str, int]:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        text = "\n".join(paragraphs)
        # Estimate page count (approx 350 words per page)
        word_count = len(text.split())
        page_count = max(1, word_count // 350)
        logger.info(f"DOCX: extracted {len(text)} chars, ~{page_count} pages")
        return text.strip(), page_count
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        raise ValueError(f"Could not extract text from DOCX: {e}")


def extract_text_from_doc(file_content: bytes) -> Tuple[str, int]:
    """Extract text from DOC (legacy Word format)."""
    try:
        # Try textract or antiword approach
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name

        try:
            import subprocess
            result = subprocess.run(
                ["antiword", tmp_path],
                capture_output=True, text=True, timeout=10
            )
            if result.returncode == 0:
                text = result.stdout
                return text.strip(), max(1, len(text.split()) // 350)
        except Exception:
            pass
        finally:
            os.unlink(tmp_path)

        raise ValueError("Legacy .doc format not supported. Please convert to .docx or .pdf")
    except Exception as e:
        raise ValueError(str(e))


def parse_resume_file(file_content: bytes, filename: str) -> Tuple[str, int]:
    """
    Main entry point for resume parsing.
    Returns: (extracted_text, page_count)
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_content)
    elif ext == ".docx":
        return extract_text_from_docx(file_content)
    elif ext == ".doc":
        return extract_text_from_doc(file_content)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Please upload PDF, DOC, or DOCX.")


def validate_file(file_content: bytes, filename: str, max_size_mb: int = 10) -> None:
    """Validate file size and type."""
    ext = Path(filename).suffix.lower()
    allowed = [".pdf", ".doc", ".docx"]

    if ext not in allowed:
        raise ValueError(f"File type '{ext}' not allowed. Use: {', '.join(allowed)}")

    size_mb = len(file_content) / (1024 * 1024)
    if size_mb > max_size_mb:
        raise ValueError(f"File too large: {size_mb:.1f}MB. Max allowed: {max_size_mb}MB")

    if len(file_content) < 100:
        raise ValueError("File appears to be empty or corrupted")
